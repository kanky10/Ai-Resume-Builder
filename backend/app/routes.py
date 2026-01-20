from fastapi import APIRouter, UploadFile, File, Depends
from sqlalchemy.orm import Session
from app.db import SessionLocal
from app.models import Resume
from app.schemas import ResumeCreate

import os
import uuid
import pdfplumber
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from openai import OpenAI


from app.models import User
from app.auth import hash_password, verify_password, create_access_token


# ------------------------
# Router setup
# ------------------------
router = APIRouter()

# ------------------------
# Directories
# ------------------------
UPLOAD_DIR = "uploads"
GENERATED_DIR = "generated"

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(GENERATED_DIR, exist_ok=True)

# ------------------------
# Database dependency
# ------------------------
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# ------------------------
# Resume text extraction
# ------------------------
def extract_text(file_path: str) -> str:
    if file_path.endswith(".pdf"):
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        return text

    if file_path.endswith(".docx"):
        doc = Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs)

    return ""

# ------------------------
# Manual resume save
# ------------------------
@router.post("/resume")
def create_resume(resume: ResumeCreate, db: Session = Depends(get_db)):
    new_resume = Resume(
        name=resume.name,
        email=resume.email,
        skills=resume.skills,
        experience=resume.experience
    )
    db.add(new_resume)
    db.commit()
    db.refresh(new_resume)
    return new_resume

# ------------------------
# Get all resumes
# ------------------------
@router.get("/resumes")
def get_resumes(db: Session = Depends(get_db)):
    return db.query(Resume).all()

# ------------------------
# Resume upload (PDF / DOCX)
# ------------------------
@router.post("/upload")
async def upload_resume(file: UploadFile = File(...)):
    file_path = os.path.join(UPLOAD_DIR, file.filename)

    with open(file_path, "wb") as f:
        f.write(await file.read())

    extracted_text = extract_text(file_path)

    return {
        "filename": file.filename,
        "extracted_text": extracted_text[:2000]
    }

# ------------------------
# ATS scoring (internal logic)
# ------------------------
@router.post("/ats-score")
def ats_score(payload: dict):
    resume_text = payload.get("resume_text", "").lower()
    keywords = payload.get("job_keywords", [])

    if not resume_text or not keywords:
        return {"score": 0}

    matched = sum(1 for kw in keywords if kw.lower() in resume_text)
    score = int((matched / len(keywords)) * 100)

    return {
        "matched_keywords": matched,
        "total_keywords": len(keywords),
        "score": score
    }

# ------------------------
# AI resume enhancement
# ------------------------
client = OpenAI()

@router.post("/ai-enhance")
def ai_enhance(payload: dict):
    resume_text = payload.get("resume_text", "")
    target_role = payload.get("target_role", "")

    prompt = f"""
    Improve the following resume for the role of {target_role}.
    Make it professional, ATS-friendly, concise, and keyword optimized.

    Resume:
    {resume_text}
    """

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are an expert resume writer."},
            {"role": "user", "content": prompt}
        ]
    )

    enhanced_text = response.choices[0].message.content

    return {
        "enhanced_resume": enhanced_text
    }

# ------------------------
# Generate PDF
# ------------------------
@router.post("/generate-pdf")
def generate_pdf(payload: dict):
    resume_text = payload.get("resume_text", "")
    file_id = str(uuid.uuid4())
    pdf_path = os.path.join(GENERATED_DIR, f"{file_id}.pdf")

    c = canvas.Canvas(pdf_path, pagesize=A4)
    text_obj = c.beginText(40, 800)

    for line in resume_text.split("\n"):
        text_obj.textLine(line)

    c.drawText(text_obj)
    c.save()

    return {
        "pdf_file": pdf_path
    }

# ------------------------
# Generate DOCX
# ------------------------
@router.post("/generate-docx")
def generate_docx(payload: dict):
    resume_text = payload.get("resume_text", "")
    file_id = str(uuid.uuid4())
    docx_path = os.path.join(GENERATED_DIR, f"{file_id}.docx")

    doc = Document()
    for line in resume_text.split("\n"):
        doc.add_paragraph(line)

    doc.save(docx_path)

    return {
        "docx_file": docx_path
    }

@router.post("/signup")
def signup(payload: dict, db: Session = Depends(get_db)):
    email = payload.get("email")
    password = payload.get("password")

    if db.query(User).filter(User.email == email).first():
        return {"error": "User already exists"}

    user = User(
        email=email,
        password=hash_password(password)
    )
    db.add(user)
    db.commit()

    return {"message": "Signup successful"}

@router.post("/login")
def login(payload: dict, db: Session = Depends(get_db)):
    email = payload.get("email")
    password = payload.get("password")

    user = db.query(User).filter(User.email == email).first()
    if not user or not verify_password(password, user.password):
        return {"error": "Invalid credentials"}

    token = create_access_token({"sub": user.email})
    return {"access_token": token}